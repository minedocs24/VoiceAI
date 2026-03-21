"""DOCX format generator - Word document with header, speaker colors, timestamps."""

from __future__ import annotations

from datetime import datetime, timezone
from io import BytesIO
from typing import Union

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.models.schemas import DiarizationResult, TranscriptResult

ExportInput = Union[TranscriptResult, DiarizationResult]


def _load_docx_config() -> dict:
    from app.core.config import load_export_config
    cfg = load_export_config()
    return cfg.get("docx", {})


def _hex_to_rgbcolor(hex_color: str) -> RGBColor:
    """Convert #RRGGBB to RGBColor."""
    hex_color = hex_color.lstrip("#")
    r, g, b = (int(hex_color[i : i + 2], 16) for i in (0, 2, 4))
    return RGBColor(r, g, b)


def _add_toc_placeholder(doc: Document) -> None:
    """Add TOC placeholder - Word will update when opened."""
    try:
        body = doc._element.body
        sect_pr = body.sectPr
        if sect_pr is None:
            sect_pr = OxmlElement("w:sdt")
        p = doc.add_paragraph()
        p.add_run("Indice (aggiornare in Word: clic destro > Aggiorna campo)")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Simple TOC field
        r = p.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        r._r.append(fldChar1)
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = 'TOC \\o "1-3" \\h \\z \\u'
        r._r.append(instr)
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "separate")
        r._r.append(fldChar2)
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "end")
        r._r.append(fldChar3)
    except Exception:
        pass  # Fallback: no TOC


class DocxGenerator:
    """Produces Word document with header, speaker sections, timestamps, footer."""

    def __init__(self):
        self._config = _load_docx_config()

    def _get_speaker_color(self, speaker: str) -> str:
        colors = self._config.get("typography", {}).get("speaker_colors", [
            "#2E75B6", "#C55A11", "#70AD47", "#5B9BD5", "#FFC000", "#7030A0"
        ])
        try:
            idx = int(speaker.replace("SPEAKER_", ""))
            return colors[idx % len(colors)]
        except (ValueError, AttributeError):
            return colors[0]

    def generate(
        self,
        data: ExportInput,
        *,
        job_id: str,
        tenant_id: str,
        project_name: str = "Trascrizione",
        duration_str: str | None = None,
        model_str: str | None = None,
    ) -> bytes:
        """Generate DOCX as bytes."""
        doc = Document()
        cfg = self._config
        typo = cfg.get("typography", {})
        font_family = typo.get("font_family", "Calibri")
        font_size = typo.get("font_size_pt", 11)

        # Header
        header_cfg = cfg.get("header", {})
        title = doc.add_heading(header_cfg.get("title", "Trascrizione"), 0)
        subtitle = header_cfg.get("subtitle_template", "Progetto: {project_name}")
        doc.add_paragraph(subtitle.format(project_name=project_name))
        meta = doc.add_paragraph()
        meta.add_run(f"Data: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
        if duration_str:
            meta.add_run(f"  |  Durata: {duration_str}")
        if model_str:
            meta.add_run(f"  |  Modello: {model_str}")
        doc.add_paragraph()

        has_diarization = isinstance(data, DiarizationResult) and any(
            s.speaker for s in data.segments
        )
        speakers_seen: set[str] = set()
        speaker_order: list[str] = []

        # Estimate page count (approx 300 words/page)
        total_words = sum(len(s.text.split()) for s in data.segments)
        add_toc = total_words > 2500 and cfg.get("toc_min_pages", 10)
        if add_toc:
            doc.add_heading("Indice", 1)
            _add_toc_placeholder(doc)
            doc.add_paragraph()

        # Content by speaker or flat
        if has_diarization:
            for seg in data.segments:
                sp = seg.speaker or "SPEAKER_00"
                if sp not in speakers_seen:
                    speakers_seen.add(sp)
                    speaker_order.append(sp)
                    doc.add_heading(sp, 2)
                p = doc.add_paragraph()
                run = p.add_run(seg.text.strip())
                run.font.size = Pt(font_size)
                run.font.name = font_family
                color = self._get_speaker_color(sp)
                run.font.color.rgb = _hex_to_rgbcolor(color)
                # Timestamp in right margin - use tab
                ts_run = p.add_run(f"\t{self._format_ts(seg.start)}")
                ts_run.font.size = Pt(9)
                ts_run.font.italic = True
        else:
            for seg in data.segments:
                p = doc.add_paragraph()
                r = p.add_run(seg.text.strip())
                r.font.size = Pt(font_size)
                r.font.name = font_family
                ts_run = p.add_run(f"\t{self._format_ts(seg.start)}")
                ts_run.font.size = Pt(9)
                ts_run.font.italic = True

        # Footer
        footer_cfg = cfg.get("footer", {})
        if footer_cfg.get("page_number") or footer_cfg.get("export_date"):
            section = doc.sections[0]
            footer = section.footer
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            parts = []
            if footer_cfg.get("export_date"):
                parts.append("Export: " + datetime.now(timezone.utc).strftime("%Y-%m-%d"))
            if footer_cfg.get("page_number"):
                parts.append("Pagina ")
            p.add_run(" | ".join(parts) if parts else "Export " + datetime.now(timezone.utc).strftime("%Y-%m-%d"))

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _format_ts(self, seconds: float) -> str:
        h = int(seconds // 3600)
        m = int((seconds % 3600) // 60)
        s = seconds % 60
        return f"{h:02d}:{m:02d}:{s:05.2f}"
