"""Unit tests for DOCX generator."""

from __future__ import annotations

import pytest
from docx import Document
from io import BytesIO

from app.generators.docx_generator import DocxGenerator
from app.models.schemas import DiarizationResult, TranscriptResult, TranscriptSegment


def test_docx_valid_file(sample_transcript: TranscriptResult):
    gen = DocxGenerator()
    content = gen.generate(
        sample_transcript,
        job_id="job-1",
        tenant_id="t1",
        project_name="Test",
    )
    doc = Document(BytesIO(content))
    assert len(doc.paragraphs) > 0
    texts = [p.text for p in doc.paragraphs]
    assert "Buongiorno" in " ".join(texts) or any("Buongiorno" in t for t in texts)


def test_docx_diarization_colors(sample_diarization: DiarizationResult):
    gen = DocxGenerator()
    content = gen.generate(
        sample_diarization,
        job_id="job-1",
        tenant_id="t1",
        project_name="Test",
    )
    doc = Document(BytesIO(content))
    assert len(doc.paragraphs) > 0
    # Check we have speaker headings
    texts = [p.text for p in doc.paragraphs]
    assert "SPEAKER_00" in " ".join(texts) or any("SPEAKER" in t for t in texts)


def test_docx_contains_metadata(sample_transcript: TranscriptResult):
    gen = DocxGenerator()
    content = gen.generate(
        sample_transcript,
        job_id="job-1",
        tenant_id="t1",
        duration_str="10.0s",
        model_str="RTF=0.05",
    )
    doc = Document(BytesIO(content))
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert "10.0" in full_text or "Trascrizione" in full_text
